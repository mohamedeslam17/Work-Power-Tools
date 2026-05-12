#!/usr/bin/env python3
"""
SEM Report Converter - Ansaldo Energia
Usage: python3 sem_convert.py vendor.pdf [output.docx]
"""
import sys, os, re, io, datetime
from pathlib import Path
import fitz
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED  = RGBColor(0xC8,0x10,0x2E)
NAVY = RGBColor(0x1A,0x1A,0x2E)
GRAY = RGBColor(0x55,0x55,0x55)
WHITE= RGBColor(0xFF,0xFF,0xFF)

# ── helpers ───────────────────────────────────────────────────────────
def _bg(c,h):
    tc=c._tc;p=tc.get_or_add_tcPr();s=OxmlElement("w:shd")
    s.set(qn("w:val"),"clear");s.set(qn("w:color"),"auto");s.set(qn("w:fill"),h);p.append(s)
def _bdr(c,color="AAAAAA",sz=4):
    tc=c._tc;p=tc.get_or_add_tcPr();b=OxmlElement("w:tcBorders")
    for side in["top","left","bottom","right"]:
        e=OxmlElement(f"w:{side}");e.set(qn("w:val"),"single")
        e.set(qn("w:sz"),str(sz));e.set(qn("w:color"),color);b.append(e)
    p.append(b)
def _nobdr(c):
    tc=c._tc;p=tc.get_or_add_tcPr();b=OxmlElement("w:tcBorders")
    for side in["top","left","bottom","right"]:e=OxmlElement(f"w:{side}");e.set(qn("w:val"),"nil");b.append(e)
    p.append(b)
def R(p,text,bold=False,size=10,color=None,italic=False):
    r=p.add_run(text);r.bold=bold;r.italic=italic
    r.font.size=Pt(size);r.font.name="Calibri"
    if color:r.font.color.rgb=color
    return r
def SP(doc,h=2):
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(h);p.paragraph_format.space_after=Pt(h)

def _set_page_size(new_sec, w_twips, h_twips, landscape=False):
    """Write w:pgSz directly into a section's sectPr with explicit twip values."""
    sectPr = new_sec._sectPr
    for old in sectPr.findall(qn('w:pgSz')):
        sectPr.remove(old)
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), str(w_twips))
    pgSz.set(qn('w:h'), str(h_twips))
    if landscape:
        pgSz.set(qn('w:orient'), 'landscape')
    sectPr.append(pgSz)

def _new_portrait_page(doc):
    new_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _set_page_size(new_sec, 11906, 16838)          # A4 portrait  21 × 29.7 cm
    new_sec.left_margin = new_sec.right_margin = Cm(1.5)
    new_sec.top_margin  = new_sec.bottom_margin = Cm(1.5)

def _new_landscape_page(doc):
    new_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _set_page_size(new_sec, 16838, 11906, landscape=True)  # A4 landscape 29.7 × 21 cm
    new_sec.left_margin = new_sec.right_margin = Cm(1.5)
    new_sec.top_margin  = new_sec.bottom_margin = Cm(1.5)

def _toc_entry(doc, label, pg):
    """TOC line with right-aligned page number and dot leader via tab stop."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '10205')  # 18 cm content width in twips
    tab.set(qn('w:leader'), 'dot')
    tabs.append(tab)
    pPr.append(tabs)
    R(p, label, size=10)
    R(p, f'\t{pg}', size=10)

def _setup_footer(section):
    """Place an auto PAGE field in the section footer, right-aligned."""
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = fp.add_run()
    r1.font.size = Pt(10); r1.font.name = 'Calibri'; r1.font.color.rgb = GRAY
    fld1 = OxmlElement('w:fldChar')
    fld1.set(qn('w:fldCharType'), 'begin')
    r1._r.append(fld1)
    r2 = fp.add_run()
    r2.font.size = Pt(10); r2.font.name = 'Calibri'; r2.font.color.rgb = GRAY
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    r2._r.append(instr)
    r3 = fp.add_run()
    r3.font.size = Pt(10); r3.font.name = 'Calibri'; r3.font.color.rgb = GRAY
    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'end')
    r3._r.append(fld2)

def _add_carbide(p, formula='M23C6', size=10):
    """Write a chemical formula with proper Word subscript for digit sequences."""
    for part in re.split(r'(\d+)', formula):
        if part.isdigit():
            r = p.add_run(part)
            r.font.subscript = True
            r.font.size = Pt(size)
            r.font.name = 'Calibri'
        elif part:
            R(p, part, size=size)

def _clean_caption(text):
    """Strip unwanted phrases and replace ASCII symbols with Unicode in caption text."""
    text = re.sub(r'\s*No indications? of[^.]*\.', '', text, flags=re.I)
    text = re.sub(r'\bsecond[- ]?phase\s+', '', text, flags=re.I)
    text = re.sub(r'\bgamma[- ]prime\b', 'γ′', text, flags=re.I)
    text = re.sub(r'\bgamma\b', 'γ', text, flags=re.I)
    return re.sub(r'\s+', ' ', text).strip()

def _R_cap(p, text, size=11, color=None, italic=False, bold=False):
    """Render caption text, subscripting digit sequences in M-type carbide formulas."""
    for part in re.split(r'(M\d+C\d*)', text):
        if re.fullmatch(r'M\d+C\d*', part):
            for seg in re.split(r'(\d+)', part):
                if not seg:
                    continue
                r = p.add_run(seg)
                if seg.isdigit():
                    r.font.subscript = True
                r.font.size = Pt(size); r.font.name = 'Calibri'
                r.bold = bold; r.italic = italic
                if color:
                    r.font.color.rgb = color
        elif part:
            R(p, part, size=size, color=color, italic=italic, bold=bold)

# ── PDF helpers ──────────────────────────────────────────────────────────
def page_text(page):
    d=page.get_text("dict");spans=[]
    for b in d["blocks"]:
        if b.get("type")!=0:continue
        for line in b.get("lines",[]):
            for s in line.get("spans",[]):
                spans.append((s["bbox"][1],s["bbox"][0],s["text"]))
    spans.sort(key=lambda s:(round(s[0]/3)*3,s[1]))
    return ' '.join(s[2] for s in spans if s[2].strip())

def caption_from_page(page,fig_num):
    d=page.get_text("dict");ph=page.rect.height;spans=[]
    cap_min=ph*0.70;cap_max=ph*0.95
    for b in d["blocks"]:
        if b.get("type")!=0:continue
        for line in b.get("lines",[]):
            for s in line.get("spans",[]):
                y=s["bbox"][1]
                if cap_min<y<cap_max and s["text"].strip():
                    spans.append((y,s["bbox"][0],s["text"]))
    spans.sort(key=lambda s:(round(s[0]/3)*3,s[1]))
    t=re.sub(r'\s+',' ',' '.join(s[2] for s in spans).strip())
    t=re.sub(r'\s*Note\s*:.*','',t);t=re.sub(r'\s*Document No.*','',t)
    if t and not re.match(r'Fig',t,re.I):t=f"Fig 1.{fig_num} shows "+t
    return t.strip()

def is_image_page(page,pdf):
    if not any(pdf.extract_image(img[0]).get('width',0)>500 for img in page.get_images()):
        return False
    d=page.get_text("dict")
    for b in d["blocks"]:
        if b.get("type")!=0:continue
        for line in b.get("lines",[]):
            lt=' '.join(s["text"] for s in line.get("spans",[]))
            if re.search(r'As-received|SEM Analysis|Location\s+[12]|Fig(?:ure)?\.?\s+1[.\s]*\d+',lt,re.I):
                return True
    return False

# ════════════════════════════════════════════════════════════
# PARSE
# ════════════════════════════════════════════════════════════
def parse(pdf_path):
    vendor=fitz.open(pdf_path)
    full='\n'.join(page_text(p) for p in vendor)

    job_m=re.search(r'Job No[:\s.]+([\d]+)',full,re.I)
    job=job_m.group(1).strip() if job_m else 'N/A'
    sn_m=re.search(r'S/N[:\s]+([A-Z0-9]+)',full,re.I)
    serial=sn_m.group(1).strip() if sn_m else 'N/A'
    mat_m=re.search(r'Material[:\s]+(IN[\w-]+)',full,re.I)
    mat=mat_m.group(1).strip() if mat_m else 'IN738'
    dt_m=re.search(r'Date[:\s]+([A-Za-z]+ \d+[a-z]*,?\s*\d{4})',full,re.I)
    date=dt_m.group(1).strip() if dt_m else '05/05/2026'

    sm=re.search(r'Job No[:\s.]+\d+\s+FR\s+(\d+)\s+(\d+)(?:nd|rd|st|th)\s+STG?\s+(BKT|BUCKET)',full,re.I)
    stage=f"MS{sm.group(1)}001 Stage {sm.group(2)} Bucket" if sm else "MS7001 Stage 3 Bucket"

    ht="Aged"
    ia="Heavy Repair"
    # Try companion _R.pdf for ht and ia
    base=str(pdf_path)
    for r_path in [base.replace('.pdf','_R.pdf'),
                   str(Path(pdf_path).parent/f"{Path(pdf_path).stem}_R.pdf")]:
        if os.path.exists(r_path):
            rpdf=fitz.open(r_path)
            rf='\n'.join(page_text(p) for p in rpdf); rpdf.close()
            ht_m=re.search(r'Heat Treatment Condition[:\s]+([^\n••]+)',rf,re.I)
            if ht_m:ht=ht_m.group(1).strip()
            ia_m=re.search(r'Incoming Assessment[:\s]+([^\n••]+)',rf,re.I)
            if ia_m:ia=ia_m.group(1).strip()
            break

    sizes=re.findall(r'measured to be ([\d.]+) microns',full,re.I)
    l1=sizes[0] if sizes else 'N/A'
    l2=sizes[1] if len(sizes)>1 else 'N/A'
    no_anom=bool(re.search(r'No (evidence|indications) of.*(needle|sigma|eta)',full,re.I))
    rts=bool(re.search(r'suitable for return to service',full,re.I))
    conclusion=''
    for r_path in [base.replace('.pdf','_R.pdf'),
                   str(Path(pdf_path).parent/f"{Path(pdf_path).stem}_R.pdf")]:
        if os.path.exists(r_path):
            rpdf2=fitz.open(r_path)
            rlast=rpdf2[-1].get_text("text")
            rpdf2.close()
            cm=re.search(r'CONCLUSION\s*\n+(.*?)(?:Location\s*\nMorphology|$)',rlast,re.DOTALL|re.I)
            if cm:conclusion=cm.group(1).strip()
            break
    if not conclusion:
        cm=re.search(r'(The metallurgical evaluation.+?NDT inspections\.)',full,re.DOTALL|re.I)
        if cm:conclusion=re.sub(r'\s+',' ',cm.group(1)).strip()

    captions={}
    for page in vendor:
        if not is_image_page(page,vendor):continue
        pt=page_text(page)
        m=re.search(r'(?:Fig|Figure)\.?\s+1[.\s]*(\d+)\s+shows',pt,re.I)
        if not m:m=re.search(r'(?:Fig|Figure)\.?\s+1[.\s]*(\d+)',pt,re.I)
        if not m:continue
        fn=m.group(1)
        if fn in captions:continue
        captions[fn]=caption_from_page(page,fn)

    for page in vendor:
        if 'As-received' not in page_text(page):continue
        d=page.get_text("dict");spans=[]
        for b in d["blocks"]:
            if b.get("type")!=0:continue
            for line in b.get("lines",[]):
                for s in line.get("spans",[]):
                    y,x,txt=s["bbox"][1],s["bbox"][0],s["text"]
                    if y>480 and x>400 and txt.strip():spans.append((y,x,txt))
        spans.sort(key=lambda s:(round(s[0]/3)*3,s[1]))
        side=re.sub(r'\s+',' ',' '.join(s[2] for s in spans).strip())
        m2=re.search(r'(Fig\s+1\.1\s+shows.+?(?:under SEM\.|SEM\.))',side,re.I|re.DOTALL)
        if m2:captions['1']=re.sub(r'\s+',' ',m2.group(1)).strip()
        break
    if '1' not in captions or not captions.get('1','').startswith('Fig'):
        captions['1']=f"Fig 1.1 shows Image of the as-received sample (ID# {job}). The specimen was first mounted, metallographically prepared and etched with Glyceregia prior to examination under SEM."

    vendor.close()
    return dict(job=job,serial=serial,material=mat,date=date,stage=stage,
                ht=ht,ia=ia,l1=l1,l2=l2,no_anom=no_anom,rts=rts,
                conclusion=conclusion,captions=captions)

# ════════════════════════════════════════════════════════════
# EXTRACT FIGURES
# ════════════════════════════════════════════════════════════
def extract_figures(pdf_path):
    doc=fitz.open(pdf_path);figs={}
    for page in doc:
        if not is_image_page(page,doc):continue
        pt=page_text(page)
        m=re.search(r'(?:Fig|Figure)\.?\s+1[.\s]*(\d+)\s+shows',pt,re.I)
        if not m:m=re.search(r'(?:Fig|Figure)\.?\s+1[.\s]*(\d+)',pt,re.I)
        if not m:continue
        fn=m.group(1)
        if fn in figs:continue

        pw=page.rect.width;ph=page.rect.height
        # Only look for header/title text in the top 30% of the page to avoid
        # mistaking image-embedded labels for the page header.
        hdr_zone = ph * 0.30

        if fn=='1':
            d1=page.get_text("dict")
            crop_rect=None
            for b in d1["blocks"]:
                if b.get("type")==1:  # image block — use its exact bbox
                    bx0,by0,bx1,by1=b["bbox"]
                    if (bx1-bx0)>pw*0.3 and (by1-by0)>ph*0.2:
                        crop_rect=fitz.Rect(bx0,by0,bx1,by1);break
            if crop_rect is None:
                hdr_bot=80.0;cap_top=ph-220.0
                for b in d1["blocks"]:
                    if b.get("type")!=0:continue
                    for line in b.get("lines",[]):
                        lt=' '.join(s["text"] for s in line.get("spans",[]))
                        bb=line["bbox"]
                        if re.search(r'As-received|SEM Analysis',lt) and bb[1]<hdr_zone:
                            if bb[3]>hdr_bot:hdr_bot=bb[3]
                        if re.match(r'\s*Fig(?:ure)?\.?\s+1\.\d+\s+shows',lt) and bb[1]>400:
                            if bb[1]<cap_top:cap_top=bb[1]
                crop_rect=fitz.Rect(18,hdr_bot+5,pw-18,cap_top-4)
            pix=page.get_pixmap(dpi=180,clip=crop_rect)
            figs[fn]={'bytes':pix.tobytes('jpeg'),'w':pix.width,'h':pix.height}
            continue

        # All other figures — prefer exact image-block bounding boxes; fall back to text-based crop
        d=page.get_text("dict")
        img_blocks=[]
        for b in d["blocks"]:
            if b.get("type")==1:
                bx0,by0,bx1,by1=b["bbox"]
                if (bx1-bx0)>pw*0.25 and (by1-by0)>ph*0.15:
                    img_blocks.append((bx0,by0,bx1,by1))

        if img_blocks:
            ibx0=min(b[0] for b in img_blocks)-4
            iby0=min(b[1] for b in img_blocks)-4
            ibx1=max(b[2] for b in img_blocks)+4
            iby1=max(b[3] for b in img_blocks)+4
            crop=fitz.Rect(max(0,ibx0),max(0,iby0),min(pw,ibx1),min(ph,iby1))
        else:
            hdr_bot=80.0;cap_top=ph-220.0
            for b in d["blocks"]:
                if b.get("type")!=0:continue
                for line in b.get("lines",[]):
                    lt=' '.join(s["text"] for s in line.get("spans",[]))
                    bb=line["bbox"]
                    if re.search(r'SEM Analysis\s*[–—-]|As-received|Location Mapping',lt) and bb[1]<hdr_zone:
                        if bb[3]>hdr_bot:hdr_bot=bb[3]
                    if re.match(r'\s*Fig(?:ure)?\s+1\.\d+\s+shows',lt) and bb[1]>400:
                        if bb[1]<cap_top:cap_top=bb[1]
            crop=fitz.Rect(18,hdr_bot+5,pw-18,cap_top-2)

        pix=page.get_pixmap(dpi=180,clip=crop)
        figs[fn]={'bytes':pix.tobytes('jpeg'),'w':pix.width,'h':pix.height}

    doc.close()
    return figs

def _default_conclusion(info):
    """Generate a template conclusion filled with parsed report metadata."""
    l1 = info.get('l1', 'N/A')
    l2 = info.get('l2', 'N/A')
    return (
        f"SEM microstructural analysis was conducted on a {info['stage']} "
        f"(S/N: {info['serial']}) manufactured from {info['material']} superalloy. "
        f"The specimen, provided in the {info['ht']}, was examined at two representative "
        f"locations under magnifications up to 10,000x.\n\n"
        f"The microstructure revealed cuboidal primary γ′ precipitates with fine secondary γ′ "
        f"within the γ matrix. Average γ′ sizes ranged between {l1} µm and {l2} µm, consistent "
        f"across the two examined locations. Stable MC carbides were observed within grains, "
        f"while M23C6 carbides were present along grain boundaries.\n\n"
        f"Based on these findings, the examined bucket (Job {info['job']}, S/N: {info['serial']}) "
        f"is considered suitable for reconditioning, subject to completion of standard NDT "
        f"inspection prior to return to service."
    )

# ════════════════════════════════════════════════════════════
# BUILD DOCX
# ════════════════════════════════════════════════════════════
def tw(cm):
    """Centimetres → OOXML twips (twentieths of a point)."""
    return round(cm * 1440 / 2.54)

def _fix_table(t, total_cm):
    """Force fixed total width + disable Word auto-fit on a table."""
    tbl = t._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for tag in (qn('w:tblW'), qn('w:tblLayout')):
        for old in tblPr.findall(tag):
            tblPr.remove(old)
    W = OxmlElement('w:tblW')
    W.set(qn('w:w'), str(tw(total_cm))); W.set(qn('w:type'), 'dxa')
    L = OxmlElement('w:tblLayout')
    L.set(qn('w:type'), 'fixed')
    tblPr.extend([W, L])

def _cantSplit(row):
    """Prevent a table row from splitting across a page break."""
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    for old in trPr.findall(qn('w:cantSplit')):
        trPr.remove(old)
    cs = OxmlElement('w:cantSplit')
    cs.set(qn('w:val'), '1')
    trPr.append(cs)

def add_two_col(doc, left_content_fn, right_bytes, right_cm=13.0, left_cm=13.5,
                caption='', img_pix=None, max_h_cm=10.0):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _fix_table(t, left_cm + right_cm)
    row = t.rows[0]
    _cantSplit(row)
    lc = row.cells[0]; lc.width = Cm(left_cm);  _nobdr(lc)
    rc = row.cells[1]; rc.width = Cm(right_cm); _nobdr(rc)
    lc._tc.get_or_add_tcPr()
    left_content_fn(lc)
    ip = rc.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    max_w = right_cm - 0.3
    if img_pix:
        w_px, h_px = img_pix
        pic_kw = dict(height=Cm(max_h_cm)) if (max_w * h_px / w_px) > max_h_cm else dict(width=Cm(max_w))
    else:
        pic_kw = dict(width=Cm(max_w))
    ip.add_run().add_picture(io.BytesIO(right_bytes), **pic_kw)
    if caption:
        clean = _clean_caption(caption)
        if clean:
            cp = rc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            cp.paragraph_format.space_before = Pt(6)
            cp.paragraph_format.keep_together = True
            _R_cap(cp, clean, size=11, color=RED, italic=True)

def build(info, figs, out_path):
    caps = info['captions']

    # Determine all SEM figure numbers (excluding cover figs 1 and 2),
    # then group into pages of 3 for a fully dynamic layout.
    sem_fig_nums = sorted([k for k in figs.keys() if k not in ('1', '2')], key=int)
    sem_chunks   = [sem_fig_nums[i:i+3] for i in range(0, len(sem_fig_nums), 3)]
    total = 4 + len(sem_chunks) + 1   # cover + TOC + intro + micro + SEM pages + summary

    today = datetime.date.today().strftime('%d %B %Y')

    doc = Document()
    sec = doc.sections[0]
    sec.page_width=Cm(21); sec.page_height=Cm(29.7)
    sec.left_margin=sec.right_margin=Cm(1.5)
    sec.top_margin=sec.bottom_margin=Cm(1.5)
    sec.header.is_linked_to_previous = False
    for p in sec.header.paragraphs: p.clear()
    _setup_footer(sec)
    doc.styles['Normal'].font.name='Calibri'; doc.styles['Normal'].font.size=Pt(10)

    # ── shared header table column widths ────────────────────────────────────────────
    WS_P  = [9.2, 4.2, 1.7, 1.6, 1.3]    # total 18.0 cm  (portrait)
    WS_LS = [13.4, 6.3, 2.4, 2.4, 2.2]   # total 26.7 cm  (landscape)

    def add_info_table(doc, page_num, landscape=False):
        ws = WS_LS if landscape else WS_P
        t=doc.add_table(rows=2,cols=5);t.style='Table Grid'
        t.alignment=WD_TABLE_ALIGNMENT.CENTER
        _fix_table(t, sum(ws))
        for j,h in enumerate(['Project / Title','Job Number.','Rev.','page','Of']):
            c=t.rows[0].cells[j];c.width=Cm(ws[j]);_bdr(c,'888888',2)
            R(c.paragraphs[0],h,size=8,color=GRAY)
        for j,(val,bold,sz) in enumerate([
            ('SEM Metallurgical Evaluation Report',False,10),(f"JC. {info['job']}",True,11),
            ('0',False,10),(str(page_num),False,10),(str(total),False,10)]):
            c=t.rows[1].cells[j];c.width=Cm(ws[j]);_bdr(c,'888888',2)
            p=c.paragraphs[0];p.paragraph_format.space_before=Pt(2);p.paragraph_format.space_after=Pt(2)
            if j==1:R(p,val,bold=True,size=sz);pp=c.add_paragraph();R(pp,info['stage'],size=8,color=GRAY)
            else:R(p,val,bold=bold,size=sz)

    def add_page_hdr(doc, page_num, landscape=False):
        add_info_table(doc, page_num, landscape=landscape)
        SP(doc,6)

    # ══ PAGE 1: COVER ════════════════════════════════════════════════════════════
    # No logo on cover — info table serves as the page header.
    add_info_table(doc, 1)

    SP(doc,12)
    t1=doc.add_paragraph(); t1.alignment=WD_ALIGN_PARAGRAPH.CENTER
    t1.paragraph_format.space_before=Pt(100)
    R(t1,'SCANNING ELECTRON MICROSCOPY',bold=True,size=20,color=NAVY)
    t2=doc.add_paragraph(); t2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    t2.paragraph_format.space_after=Pt(220)
    R(t2,'METALLURGICAL EVALUATION REPORT',bold=True,size=14,color=GRAY)

    SW=[3.2, 5.1, 5.1, 2.6]   # total 16.0 cm, centred on 18 cm portrait
    sig=doc.add_table(rows=3,cols=4); sig.style='Table Grid'
    sig.alignment=WD_TABLE_ALIGNMENT.CENTER
    _fix_table(sig, sum(SW))
    for ri,row in enumerate([['','Name','Title','Date'],
        ['Submitted','Eslam Abdelmawla','Materials Engineer', today],
        ['Approved','Khemichi Badri','Sr. Materials Engineer', today]]):
        for ci,val in enumerate(row):
            c=sig.rows[ri].cells[ci];c.width=Cm(SW[ci]);_bdr(c)
            if ri==0:_bg(c,'F0F0F0')
            R(c.paragraphs[0],val,bold=(ri==0),size=9)

    _new_portrait_page(doc)

    # ══ PAGE 2: TOC ═══════════════════════════════════════════════════════════════
    add_page_hdr(doc, 2)
    SP(doc,4)
    h=doc.add_paragraph();h.paragraph_format.space_before=Pt(6);h.paragraph_format.space_after=Pt(10)
    R(h,'TABLE OF CONTENTS',bold=True,size=13,color=NAVY)
    # Horizontal rule under heading
    pPr=h._p.get_or_add_pPr();pBdr=OxmlElement('w:pBdr')
    bot=OxmlElement('w:bottom');bot.set(qn('w:val'),'single');bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1');bot.set(qn('w:color'),'C8102E');pBdr.append(bot);pPr.append(pBdr)

    summary_page = str(total)
    for label,pg in [('TABLE OF CONTENTS','2'),('INTRODUCTION','3'),('RECAPITULATION','3'),
                     ('MICROSTRUCTURE ANALYSIS','4'),
                     ('SUMMARY OF γ′ PRECIPITATE MEASUREMENTS', summary_page),
                     ('CONCLUSION', summary_page)]:
        _toc_entry(doc, label, pg)

    _new_landscape_page(doc)

    # ══ PAGE 3: INTRO + RECAP (left) | FIG 1.1 (right) ══════
    add_page_hdr(doc,3,landscape=True)

    def left_p3(cell):
        p=cell.add_paragraph(); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(8)
        R(p,'INTRODUCTION',bold=True,size=12,color=NAVY)
        p2=cell.add_paragraph(); p2.paragraph_format.space_after=Pt(10)
        p2.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        R(p2,f"This report presents the metallurgical evaluation of a {info['stage']} using "
            f"Scanning Electron Microscopy (SEM). The analysis was performed on the specimen in the "
            f"{info['ht']} condition. The objective is to evaluate microstructural integrity, "
            f"focusing on the γ′ morphology and the presence of any degradation phases such as "
            f"brittle needle-shaped precipitates.",size=11)
        p3=cell.add_paragraph(); p3.paragraph_format.space_before=Pt(10); p3.paragraph_format.space_after=Pt(8)
        R(p3,'RECAPITULATION',bold=True,size=12,color=NAVY)
        for lbl,val in [('Job Number',info['job']),('Alloy',info['material']),
                        ('Incoming Assessment',info['ia']),
                        ('Heat Treatment Condition',info['ht']),('Serial Number',info['serial'])]:
            pb=cell.add_paragraph(); pb.paragraph_format.space_before=Pt(2); pb.paragraph_format.space_after=Pt(2)
            pb.paragraph_format.left_indent=Cm(0.3); pb.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            R(pb,'• ',bold=True,size=11); R(pb,lbl+': ',bold=True,size=11); R(pb,val,size=11)

    if '1' in figs:
        f1=figs['1']
        add_two_col(doc,left_p3,f1['bytes'],right_cm=12.2,left_cm=14.5,
                    caption=caps.get('1','Fig 1.1'),img_pix=(f1['w'],f1['h']))
    else:
        left_p3_para=doc.add_paragraph(); left_p3(left_p3_para)

    _new_landscape_page(doc)

    # ══ PAGE 4: MICROSTRUCTURE (left) | FIG 1.2 (right) ═════
    add_page_hdr(doc,4,landscape=True)

    def left_p4(cell):
        p=cell.add_paragraph();p.paragraph_format.space_before=Pt(4);p.paragraph_format.space_after=Pt(8)
        R(p,'MICROSTRUCTURE ANALYSIS',bold=True,size=12,color=NAVY)
        p2=cell.add_paragraph();p2.paragraph_format.space_after=Pt(8)
        p2.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        R(p2,'The analysis focused on two representative locations, revealing a matrix of '
             'γ′ precipitates and various carbide phases.',size=11)
        pb=cell.add_paragraph();pb.paragraph_format.left_indent=Cm(0.3);pb.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        R(pb,'• ',size=11);R(pb,'γ Matrix: ',bold=True,size=11)
        R(pb,'Both locations showed a typical distribution of primary and secondary γ′ precipitates.',size=11)
        pb=cell.add_paragraph();pb.paragraph_format.left_indent=Cm(0.3);pb.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        R(pb,'• ',size=11);R(pb,'Precipitates:',bold=True,size=11)
        pb=cell.add_paragraph();pb.paragraph_format.left_indent=Cm(0.8);pb.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        R(pb,'o ',size=11);R(pb,'Grain Boundaries: ',bold=True,size=11)
        R(pb,'Fine and coarse precipitates, identified as likely ',size=11)
        _add_carbide(pb,size=11)
        R(pb,' and MC-type carbides, were observed along the grain boundaries.',size=11)
        pb=cell.add_paragraph();pb.paragraph_format.left_indent=Cm(0.8);pb.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        R(pb,'o ',size=11);R(pb,'Intra-granular: ',bold=True,size=11)
        R(pb,'Coarse, blocky MC-type precipitates were found within the grains.',size=11)

    if '2' in figs:
        f2=figs['2']
        add_two_col(doc,left_p4,f2['bytes'],right_cm=13.0,left_cm=13.5,
                    caption=caps.get('2','Fig 1.2'),img_pix=(f2['w'],f2['h']))
    _new_landscape_page(doc)

    # ══ SEM IMAGE PAGES: 3 figures per page, location label from captions ═══
    def _loc_label(fn_list):
        """Detect location from captions; fall back to 'Location 1'."""
        for fn in fn_list:
            cap = caps.get(fn, '').lower()
            if 'location 2' in cap:
                return 'Location 2'
            elif 'location 1' in cap:
                return 'Location 1'
        return 'Location 1'

    def sem_page(nums, loc_lbl, page_num, next_portrait=False):
        add_page_hdr(doc, page_num, landscape=True)
        present=[(n,figs[n]) for n in nums if n in figs]
        if not present:return
        col_cm=8.85;img_cm=8.55;max_h=8.5

        t=doc.add_table(rows=1,cols=3);t.alignment=WD_TABLE_ALIGNMENT.CENTER
        _fix_table(t,3*col_cm)
        _cantSplit(t.rows[0])
        for ci in range(3):
            cell=t.rows[0].cells[ci];cell.width=Cm(col_cm);_nobdr(cell)
            if ci>=len(present):continue
            fn,f=present[ci]
            ip=cell.paragraphs[0];ip.alignment=WD_ALIGN_PARAGRAPH.CENTER
            ip.paragraph_format.space_before=Pt(6)
            w_px,h_px=f['w'],f['h']
            pic_kw=dict(height=Cm(max_h)) if (img_cm*h_px/w_px)>max_h else dict(width=Cm(img_cm))
            ip.add_run().add_picture(io.BytesIO(f['bytes']),**pic_kw)
            cp=cell.add_paragraph();cp.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            cp.paragraph_format.space_before=Pt(5)
            cp.paragraph_format.space_after=Pt(8)
            cp.paragraph_format.keep_together=True
            _R_cap(cp,_clean_caption(caps.get(fn,f'Fig 1.{fn}')),size=12,color=GRAY,italic=True)

        # Location label as a merged row fixed inside the table
        label_row=t.add_row()
        _cantSplit(label_row)
        lc=label_row.cells[0].merge(label_row.cells[2])
        _nobdr(lc)
        ll_p=lc.paragraphs[0];ll_p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        ll_p.paragraph_format.space_before=Pt(8);ll_p.paragraph_format.space_after=Pt(4)
        R(ll_p,loc_lbl,bold=True,size=12)

        if next_portrait:
            _new_portrait_page(doc)
        else:
            _new_landscape_page(doc)

    for i, chunk in enumerate(sem_chunks):
        page_num  = 5 + i
        is_last   = (i == len(sem_chunks) - 1)
        sem_page(chunk, _loc_label(chunk), page_num, next_portrait=is_last)

    # ══ SUMMARY + CONCLUSION ════════════════════════════════════════════════
    add_page_hdr(doc, total)

    h=doc.add_paragraph();h.paragraph_format.space_before=Pt(10);h.paragraph_format.space_after=Pt(6)
    R(h,'SUMMARY OF γ′ PRECIPITATE MEASUREMENTS',bold=True,size=11,color=NAVY)
    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(8)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    R(p,'To consolidate the observations from both examined locations, the measured sizes of primary γ′ precipitates are summarized in the table below:')
    SP(doc,4)

    gt=doc.add_table(rows=3,cols=3);gt.alignment=WD_TABLE_ALIGNMENT.CENTER;gt.style='Table Grid'
    gw=[Cm(5.0),Cm(5.0),Cm(5.0)]
    for j,h in enumerate(['Location','Morphology','Average Size (μm)']):
        c=gt.rows[0].cells[j];c.width=gw[j];_bg(c,'1A1A2E');_bdr(c)
        cp=c.paragraphs[0];cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        R(cp,h,bold=True,size=10,color=WHITE)
    for ri,row in enumerate([['Location 1','Cuboidal',info['l1']],['Location 2','Cuboidal',info['l2']]]):
        for ci,val in enumerate(row):
            c=gt.rows[ri+1].cells[ci];c.width=gw[ci];_bdr(c)
            if ri==1:_bg(c,'F8F8F8')
            cp=c.paragraphs[0];cp.alignment=WD_ALIGN_PARAGRAPH.CENTER;R(cp,val,size=10)
    tp=doc.add_paragraph();tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before=Pt(6);tp.paragraph_format.space_after=Pt(12)
    R(tp,'Table 1 Primary γ′ Precipitate Size Measurements',italic=True,size=9,color=GRAY)

    h=doc.add_paragraph();h.paragraph_format.space_before=Pt(10);h.paragraph_format.space_after=Pt(6)
    R(h,'CONCLUSION',bold=True,size=11,color=NAVY)
    conclusion_text=info['conclusion'] if info['conclusion'] else _default_conclusion(info)
    for part in re.split(r'\n\s*\n',conclusion_text):
        clean=re.sub(r'\s*No indications? of[^.]*\.','',part,flags=re.I)
        clean=re.sub(r'\s+',' ',clean).strip()
        if clean:
            p=doc.add_paragraph();p.paragraph_format.space_after=Pt(8)
            p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            _R_cap(p,clean,size=10)

    doc.save(out_path)

# ════════════════════════════════════════════════════════════
def main():
    if len(sys.argv)<2:print(__doc__);sys.exit(1)
    pdf=sys.argv[1]
    if not os.path.exists(pdf):sys.exit(f'Not found: {pdf}')
    out=sys.argv[2] if len(sys.argv)>2 else f'Ansaldo_{Path(pdf).stem}.docx'
    print(f'Converting: {pdf}')
    print('  → Parsing...')
    info=parse(pdf)
    print(f"     {info['stage']}  HT: {info['ht']}")
    print(f"     γ′: L1={info['l1']} L2={info['l2']}")
    print('  → Extracting figures...')
    figs=extract_figures(pdf)
    print(f'     {len(figs)} figures: {sorted(figs.keys(),key=int)}')
    print('  → Building Word document...')
    build(info,figs,out)
    print(f'\n✅  Done: {out}  ({os.path.getsize(out)//1024} KB)')

if __name__=='__main__':main()
